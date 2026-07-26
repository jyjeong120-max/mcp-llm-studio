"""files.py

첨부 파일을 uploads/에 저장하고 텍스트를 추출한다.

PDF/DOCX 추출은 선택 의존성(pypdf, python-docx)이 있을 때만 동작하고,
없으면 해당 형식만 안내 메시지로 대체한다(우아한 저하 — 앱은 계속 동작).
"""

from __future__ import annotations

import json
import uuid
from pathlib import Path

MAX_FILE_BYTES = 50 * 1024 * 1024   # 업로드 원본 한도
MAX_EXTRACT_CHARS = 200_000         # 추출 텍스트 한도 (컨텍스트 보호)

TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".py", ".js", ".ts", ".html", ".css",
    ".xml", ".yaml", ".yml", ".ini", ".log", ".sql", ".java", ".c", ".cpp", ".sh", ".bat",
}

# 서버가 텍스트로 못 뽑는 바이너리 오피스 포맷. 추출 대신 저장 경로를 모델에 주고,
# 모델이 확장자에 맞는 MCP 도구(office 등)를 골라 그 경로를 읽게 한다(path 모드).
# xlsx/pptx는 애초에 서버 추출기가 없으므로 스위치와 무관하게 항상 path 모드.
PATH_EXTENSIONS = {
    ".xlsx", ".xls", ".xlsm", ".xlsb",   # Excel
    ".pptx", ".ppt", ".pptm",            # PowerPoint
}

# 사내 DRM 스위치(attachment_com_office)가 켜졌을 때 path 모드로 돌릴 Office 포맷 전체.
# docx는 평소엔 python-docx로 추출하지만, DRM 암호화 환경에선 바이트가 암호문이라
# 추출이 실패하므로 이때만 COM 경로 읽기로 넘긴다.
OFFICE_COM_EXTENSIONS = PATH_EXTENSIONS | {".docx", ".doc", ".docm"}


class UploadStore:
    def __init__(self, base_dir: Path):
        self.dir = base_dir / "uploads"
        self.dir.mkdir(exist_ok=True)

    def save(self, filename: str, data: bytes, com_office: bool = False) -> dict:
        """원본을 저장하고, 형식에 따라 텍스트 추출(text) 또는 경로 전달(path) 모드로 메타를 만든다.

        com_office=True면 DRM 환경으로 보고 docx 등 Office 문서도 추출하지 않고
        경로 모드로 넘긴다(모델이 office COM 도구로 열어 복호화하며 읽게).
        """
        if len(data) > MAX_FILE_BYTES:
            raise ValueError(f"파일이 너무 큽니다 (한도 {MAX_FILE_BYTES // (1024*1024)}MB)")
        file_id = uuid.uuid4().hex[:12]
        safe_name = Path(filename).name
        original = self.dir / f"{file_id}_{safe_name}"
        original.write_bytes(data)

        ext = Path(safe_name).suffix.lower()
        path_mode = ext in PATH_EXTENSIONS or (com_office and ext in OFFICE_COM_EXTENSIONS)
        if path_mode:
            # 텍스트로 못 뽑는(또는 DRM으로 못 뽑는) 포맷: 경로 모드. 내용은 모델이 도구로 읽는다.
            mode, text = "path", ""
        else:
            mode, text = "text", extract_text(safe_name, data)
        meta = {
            "id": file_id,
            "name": safe_name,
            "chars": len(text),
            "mode": mode,
            # 백슬래시는 모델이 JSON 인자로 넘길 때 깨지기 쉬워 슬래시로 준다 (COM도 허용).
            "path": str(original.resolve()).replace("\\", "/"),
        }
        (self.dir / f"{file_id}.txt").write_text(text, encoding="utf-8")
        (self.dir / f"{file_id}.meta.json").write_text(
            json.dumps(meta, ensure_ascii=False), encoding="utf-8"
        )
        return meta

    def register_path(self, abs_path: str, com_office: bool = False) -> dict:
        """디스크의 원본 파일을 '복사 없이' 경로 참조로 등록한다.

        확장자로 처리 방식을 가른다(save와 동일 규칙):
        - Office 바이너리(xlsx/pptx) 또는 com_office=True인 DRM 대상(docx 등)
          → path 모드. 사본을 만들지 않고 원본 경로만 남겨, office COM 도구가
          원본을 제자리에서 열어(복호화하며) 읽게 한다.
          (DRM 환경에서 업로드=사본 생성은 암호화 봉투를 상하게 해 복호화가 깨진다.)
        - 그 외(txt/pdf/일반 docx 등) → 원본 바이트를 '읽어서'(복사 아님) 텍스트를
          추출해 인라인한다. 붙는 도구가 없어도 내용을 볼 수 있게. 읽기 자체는
          DRM 봉투를 건드리지 않는다.
        - 지원 확장자가 아니거나 읽기/추출이 실패하면 path 모드로 물러선다.
        """
        p = Path(abs_path)
        if not p.is_file():
            raise ValueError(f"파일을 찾을 수 없습니다: {abs_path}")
        try:
            size = p.stat().st_size
        except OSError as e:
            raise ValueError(f"파일 정보를 읽을 수 없습니다: {e}")
        if size > MAX_FILE_BYTES:
            raise ValueError(f"파일이 너무 큽니다 (한도 {MAX_FILE_BYTES // (1024*1024)}MB)")

        file_id = uuid.uuid4().hex[:12]
        ext = p.suffix.lower()
        path_mode = ext in PATH_EXTENSIONS or (com_office and ext in OFFICE_COM_EXTENSIONS)
        if path_mode:
            mode, text = "path", ""
        elif ext in TEXT_EXTENSIONS or ext in (".pdf", ".docx"):
            # 서버가 뽑을 수 있는 형식: 원본을 읽어 텍스트로. 실패하면 경로 모드로 저하.
            try:
                mode, text = "text", extract_text(p.name, p.read_bytes())
            except Exception:  # noqa: BLE001 — 권한/손상/DRM 등
                mode, text = "path", ""
        else:
            # png/zip 등 애초에 텍스트가 없는 형식: 인라인도 도구 오지시도 하지 않도록
            # path 모드로 두되, get() 텍스트는 비어 모델이 '읽을 수 없음'을 알게 한다.
            mode, text = "path", ""
        meta = {
            "id": file_id,
            "name": p.name,
            "chars": len(text),
            "mode": mode,
            # 원본 절대경로(슬래시). 사본이 아니라 제자리 원본을 가리킨다.
            "path": str(p.resolve()).replace("\\", "/"),
            "ref": True,  # 업로드 사본이 아닌 원본 경로 참조 표시
        }
        if text:
            (self.dir / f"{file_id}.txt").write_text(text, encoding="utf-8")
        (self.dir / f"{file_id}.meta.json").write_text(
            json.dumps(meta, ensure_ascii=False), encoding="utf-8"
        )
        return meta

    def get(self, file_id: str) -> tuple[dict, str] | None:
        """(메타, 추출 텍스트)를 돌려준다. 없으면 None. path 모드면 텍스트는 빈 문자열."""
        safe = "".join(ch for ch in file_id if ch.isalnum())
        meta_path = self.dir / f"{safe}.meta.json"
        if not meta_path.exists():
            return None
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        text_path = self.dir / f"{safe}.txt"
        text = text_path.read_text(encoding="utf-8") if text_path.exists() else ""
        return meta, text


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in TEXT_EXTENSIONS:
        text = _decode(data)
    elif ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    else:
        text = f"[지원하지 않는 형식입니다: {ext} — 텍스트/PDF/DOCX 파일을 사용하세요]"
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + f"\n\n[이하 생략 — 총 {len(text):,}자 중 앞부분만 포함]"
    return text


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "cp949", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        import io
        from pypdf import PdfReader
    except ImportError:
        return "[PDF 추출을 쓰려면 pypdf 패키지가 필요합니다 (pip install pypdf)]"
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(pages)
    except Exception as e:  # noqa: BLE001
        return f"[PDF 추출 실패: {e}]"


def _extract_docx(data: bytes) -> str:
    try:
        import io
        from docx import Document
    except ImportError:
        return "[DOCX 추출을 쓰려면 python-docx 패키지가 필요합니다 (pip install python-docx)]"
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:  # noqa: BLE001
        return f"[DOCX 추출 실패: {e}]"
